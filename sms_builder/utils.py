import re
import shutil
import os
import subprocess
import tempfile
import time

from django.conf import settings
from django.core.files import File
from django.http import Http404
from docx import Document

from .models import CompanyDocument






def format_step1_data(data):
    """Validates and formats Company Details"""
    # Clean the ABN (remove spaces)
    abn_raw = data.get('abn', '')
    abn_clean = re.sub(r'\s+', '', str(abn_raw)) if abn_raw else None

    return {
        'company_name': data.get('company_name'),
        'abn': abn_clean,
        'address_street': data.get('address_street'),
        'city': data.get('city'),
        'state': data.get('state'),
        'postcode': data.get('postcode'),
        'contact_phone': data.get('contact_phone'),
        'contact_email': data.get('contact_email'),
        'contact_person': data.get('contact_person'),
        'contact_role': data.get('contact_role'),
        'declaration_accepted': data.get('declaration_accepted', False)
    }

def format_step2_data(data):
    """Validates and formats Operations Details"""
    
    # Helper to convert JS empty date strings "" to None for Django DateFields
    def parse_date(date_str):
        return date_str if date_str and date_str.strip() != "" else None

    # Handle numeric fallbacks
    try:
        num_drivers = int(data.get('num_drivers', 0))
    except (ValueError, TypeError):
        num_drivers = 0

    return {
        'work_types': data.get('work_types', []),
        'accreditations': data.get('accreditations', []),
        'audit_date_none': parse_date(data.get('audit_date_none')),
        'audit_date_trucksafe': parse_date(data.get('audit_date_trucksafe')),
        'audit_date_wahva': parse_date(data.get('audit_date_wahva')),
        'operating_areas': data.get('operating_areas', []),
        'operating_hours': data.get('operating_hours'),
        'num_drivers': num_drivers,
    }


def format_step3_data(data):
    """Validates and formats Fleet Details"""
    
    # Handle numeric fallbacks for vehicles
    try:
        total_vehicles = int(data.get('totalVehicles', 0))
    except (ValueError, TypeError):
        total_vehicles = 0

    return {
        'total_vehicles': total_vehicles,
        'max_gvm': data.get('maxGVM'),
        'average_vehicle_age': data.get('vehicleAge'),
        'vehicle_types': data.get('vehicle_types', []),
        'special_cargo': data.get('special_cargo', []),
        # Ensure it defaults to an empty dict as per your model
        'nhvr_configurations': data.get('nhvr_configurations', {}) 
    }

def format_step4_static_data(data):
    """Validates and formats Static Risk Profile Details"""
    return {
        'safety_policies': data.get('safety_policies', []),
        'additional_notes': data.get('riskNotes', '')
    }

def format_step4_dynamic_data(data):
    """Extracts dynamic risk hazards from the payload"""
    return data.get('risk_hazards', [])    


def format_step5_static_data(data):
    """Validates and formats Static Subcontractor Details"""
    def parse_int(val):
        try:
            return int(val) if val and str(val).strip() != "" else 0
        except (ValueError, TypeError):
            return 0

    return {
        'engages_subcontractors': data.get('engages_subcontractors', False),
        'compliance_practices': data.get('compliance_practices', []),
        'active_subcontractors': parse_int(data.get('active_subcontractors')),
        'primary_engagement_type': data.get('primary_engagement_type', ''),
        'review_frequency': data.get('review_frequency', ''),
        'cor_procedures': data.get('cor_procedures', '')
    }

def format_step5_dynamic_data(data):
    """Extracts dynamic subcontractor records"""
    return data.get('subcontractor_records', [])

def format_step6_static_data(data):
    """Validates and formats Static Incident Details"""
    def parse_int(val):
        try:
            return int(val) if val and str(val).strip() != "" else 0
        except (ValueError, TypeError):
            return 0

    return {
        'reporting_process': data.get('reporting_process', []),
        'incidents_last_12_months': parse_int(data.get('incidents_last_12_months')),
        'incidents_last_3_years': parse_int(data.get('incidents_last_3_years')),
        'injuries_resulting': parse_int(data.get('injuries_resulting')),
        'improvement_actions': data.get('improvement_actions', '')
    }

def format_step6_dynamic_data(data):
    """Extracts dynamic incident records"""
    return data.get('incident_records', [])


# def get_libreoffice():
#     if os.name == "nt":
#         possible_paths = [
#             r"C:\Program Files\LibreOffice\program\soffice.exe",
#             r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
#         ]
#         for path in possible_paths:
#             if os.path.exists(path):
#                 return path

#     return shutil.which("soffice") or shutil.which("libreoffice")


def generate_company_document_for_company(company):
    """
    Generates the company master document and saves it as FULL_DOC.
    Returns the saved CompanyDocument object.
    """

    # Template path
    template_path = os.path.join(
        settings.BASE_DIR,
        'sms_builder',
        'assets',
        'master_template.docx'
    )

    if not os.path.exists(template_path):
        raise Http404("Master template document not found.")

    # Load DOCX
    doc = Document(template_path)

    # Build address
    address_parts = [
        company.address_street,
        company.city,
        company.state,
        company.postcode
    ]
    business_location = ", ".join(filter(None, address_parts))

    if not business_location:
        business_location = company.address or "Location not provided"

    # Replace placeholders
    replacements = {
        "[INSERT COMPANY NAME]": company.company_name or "Company Name Missing",
        "[INSERT BUSINESS LOCATION]": business_location,
    }

    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in p.text:
                            p.text = p.text.replace(old, new)

    # Temp paths
    temp_dir = tempfile.gettempdir()
    docx_path = os.path.join(temp_dir, f"{company.id}_master.docx")
    pdf_path = os.path.join(temp_dir, f"{company.id}_master.pdf")

    doc.save(docx_path)

    # Get LibreOffice
    libreoffice = get_libreoffice()
    if not libreoffice:
        raise Exception("LibreOffice not found on system")

    # Convert DOCX → PDF
    subprocess.run([
        libreoffice,
        "--headless",
        "--convert-to", "pdf",
        docx_path,
        "--outdir", temp_dir
    ], check=True)

    # Wait for PDF
    for _ in range(10):
        if os.path.exists(pdf_path):
            break
        time.sleep(0.5)

    if not os.path.exists(pdf_path):
        raise Http404("PDF generation failed.")

    # Save to model
    safe_name = (company.company_name or "Company").replace(" ", "_")
    file_name = f"{safe_name}_SMS_Document.pdf"

    with open(pdf_path, 'rb') as f:
        CompanyDocument.objects.filter(
            company=company,
            doc_type="FULL_DOC"
        ).delete()

        doc_obj = CompanyDocument.objects.create(
            company=company,
            file=File(f, name=file_name),
            name=file_name,
            doc_type="FULL_DOC"
        )

    # Cleanup temp files
    try:
        os.remove(docx_path)
        os.remove(pdf_path)
    except:
        pass

    return doc_obj

def get_libreoffice():
    return shutil.which("soffice")

